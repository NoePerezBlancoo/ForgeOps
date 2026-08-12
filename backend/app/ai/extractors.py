from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader


class UnsupportedDocumentError(ValueError):
    pass


class EmptyDocumentError(ValueError):
    pass


@dataclass(frozen=True)
class ExtractedSection:
    text: str
    page_number: int | None = None


def extract_sections(content: bytes, original_name: str, max_chars: int) -> list[ExtractedSection]:
    extension = Path(original_name).suffix.lower()
    if extension == ".txt":
        sections = _extract_text(content)
    elif extension == ".pdf":
        sections = _extract_pdf(content)
    elif extension == ".docx":
        sections = _extract_docx(content)
    else:
        raise UnsupportedDocumentError(
            "El formato no dispone de extraccion de texto. Usa TXT, PDF o DOCX."
        )

    clean_sections = [
        ExtractedSection(text=_normalize(section.text), page_number=section.page_number)
        for section in sections
        if _normalize(section.text)
    ]
    if not clean_sections:
        raise EmptyDocumentError("El documento no contiene texto extraible")
    if sum(len(section.text) for section in clean_sections) > max_chars:
        raise ValueError("El documento supera el limite de texto indexable")
    return clean_sections


def _extract_text(content: bytes) -> list[ExtractedSection]:
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            return [ExtractedSection(content.decode(encoding))]
        except UnicodeDecodeError:
            continue
    raise EmptyDocumentError("No se pudo decodificar el archivo de texto")


def _extract_pdf(content: bytes) -> list[ExtractedSection]:
    reader = PdfReader(BytesIO(content))
    if reader.is_encrypted:
        try:
            reader.decrypt("")
        except Exception as exc:
            raise UnsupportedDocumentError("El PDF esta protegido con contrasena") from exc
    return [
        ExtractedSection(text=page.extract_text() or "", page_number=index + 1)
        for index, page in enumerate(reader.pages)
    ]


def _extract_docx(content: bytes) -> list[ExtractedSection]:
    document = Document(BytesIO(content))
    blocks = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                blocks.append(" | ".join(values))
    return [ExtractedSection("\n\n".join(blocks))]


def _normalize(value: str) -> str:
    lines = [" ".join(line.replace("\x00", "").split()) for line in value.splitlines()]
    return "\n".join(line for line in lines if line).strip()

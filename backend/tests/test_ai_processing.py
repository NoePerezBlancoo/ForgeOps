from docx import Document

from app.ai.chunking import chunk_sections
from app.ai.extractors import ExtractedSection, extract_sections


def test_docx_extraction_includes_paragraphs_and_tables(tmp_path):
    path = tmp_path / "manual.docx"
    document = Document()
    document.add_paragraph("Procedimiento de seguridad   para mantenimiento")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Paso"
    table.cell(0, 1).text = "Aislar energia"
    document.save(path)

    sections = extract_sections(path, path.name, max_chars=10_000)

    assert len(sections) == 1
    assert "Procedimiento de seguridad para mantenimiento" in sections[0].text
    assert "Paso | Aislar energia" in sections[0].text


def test_chunking_preserves_page_and_context_overlap():
    content = " ".join(f"termino{index:03d}" for index in range(120))

    chunks = chunk_sections(
        [ExtractedSection(text=content, page_number=4)],
        target_chars=320,
        overlap_chars=80,
    )

    assert len(chunks) > 2
    assert all(chunk.page_number == 4 for chunk in chunks)
    assert all(chunk.token_count > 0 for chunk in chunks)
    for previous, current in zip(chunks[:-1], chunks[1:], strict=True):
        previous_tail = set(previous.content.split()[-8:])
        current_head = set(current.content.split()[:12])
        assert previous_tail & current_head
